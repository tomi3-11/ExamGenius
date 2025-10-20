# importing the necessary tools from the flask framework
from flask import Flask, request, jsonify, send_file, render_template
from fpdf import FPDF
import io
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER

# Create an instance of the Flask Application
app = Flask(__name__)
CORS(app)


# Custom PDF class using Unicode safe fonts
class PDF(FPDF):
    pass

# Register Unicode font
FONT_PATH = "fonts/DejaVuSans.ttf"

# Defining a route for the pdf application
# Route for the Homepage
@app.route('/')
def home():
    """Serves the main HTML page."""
    return render_template('index.html')

# PDF Generation Logic
@app.route('/generate-pdf', methods=['POST'])
def generate_pdf():
    """
    Receives exam data from the frontend, creates a pdf,
    and sends for download.
    """
    try:
        # Get the JSON data sent from the Frontend
        data = request.get_json()
        
        title = data.get('title', 'Exam Paper')
        instructions = data.get('instructions', '')
        questions = data.get('questions', [])
        
        pdf = FPDF() 
        pdf.add_page()
        
        pdf.add_font('DejaVu', "", FONT_PATH) 
        pdf.add_font("DejaVu", "B", FONT_PATH) 
        
        # adding Title
        pdf.set_font('DejaVu', "B", 16) 
        pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT", align='C')
        
        # Add instructions
        if instructions:
            pdf.set_font('DejaVu', '', 12)
            pdf.multi_cell(0, 10, instructions)
            pdf.ln(10)
            
        # Add questions
        pdf.set_font('DejaVu', '', 12)
        for i, question_data in enumerate(questions, 1):
            
            q_text = question_data.get('text', '').replace("→", "->")
            answer_lines = question_data.get('lines', 0)
            is_blank_only = question_data.get('blankOnly', False)
            
            q_type = question_data.get('type', 'short_answer')
            
            pdf.multi_cell(0, 10, f'{i}. {q_text}')
            
            if q_type == 'mcq':
                options = question_data.get('options', [])
                pdf.ln(5)
                
                letter_cell_width = 15
                
                text_cell_width = pdf.epw - letter_cell_width
                
                for j, option_text in enumerate(options):
                    letter = chr(97 + j)
                    pdf.cell(letter_cell_width, 10, f'     {letter}.)', align='L')
                    pdf.multi_cell(text_cell_width, 10, option_text, new_x="LMARGIN", new_y="NEXT")
                    
                pdf.ln(7)
                
            else:
                if answer_lines > 0:
                    pdf.ln(5) 
                    if not is_blank_only:
                        pdf.set_font('DejaVu', '', 11) 
                        pdf.cell(0, 10, "Answer:")
                        pdf.ln(5)
                    
                    pdf.set_font('DejaVu', '', 12)
                    for _ in range(answer_lines):
                        pdf.cell(0, 10, "", border='B', new_x="LMARGIN", new_y="NEXT") 
                    pdf.ln(7)
                else:
                    pdf.ln(5)
            
        pdf_bytes = pdf.output() 
        pdf_buffer = io.BytesIO(pdf_bytes)
        
        # Send the PDF back to the browser
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name='exam.pdf',
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"Error occured (PDF): {e}")
        return jsonify({"error": "Failed to generate PDF"}), 500


@app.route('/generate-docx', methods=['POST'])
def generate_docx():
    """
    Receives exam data, creates a DOCX, and sends it back for download.
    """
    try:
        data = request.get_json()
        
        title = data.get('title', 'Exam Paper')
        instructions = data.get('instructions', '')
        questions = data.get('questions', [])

        doc = Document()

        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        doc.add_paragraph()

        if instructions:
            doc.add_paragraph(instructions)
            doc.add_paragraph()

        section = doc.sections[0]
        tab_position = section.page_width - section.left_margin - section.right_margin
        
        for i, question_data in enumerate(questions, 1):
            
            q_type = question_data.get('type', 'short_answer')
            q_text = question_data.get('text', '')
            answer_lines = question_data.get('lines', 0)
            is_blank_only = question_data.get('blankOnly', False)
            
            doc.add_paragraph(f'{i}. {q_text}')
            
            if q_type == 'mcq':
                options = question_data.get('options', [])
                
                for j, option_text in enumerate(options):
                    letter = chr(97 + j)
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.add_run(f'{letter}) {option_text}')
                    
                doc.add_paragraph()

            if answer_lines > 0:
                if not is_blank_only:
                    answer_prompt_p = doc.add_paragraph()
                    answer_run = answer_prompt_p.add_run("Answer:")
                    answer_run.font.italic = True
                
                for _ in range(answer_lines):
                    p = doc.add_paragraph()
                    p_format = p.paragraph_format
                    tab_stops = p_format.tab_stops
                    tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.HEAVY)
                    p.add_run('\t')
                    
                doc.add_paragraph()
            else:
                doc.add_paragraph()

        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name='exam.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"An error occurred (DOCX): {e}")
        return jsonify({"error": "Failed to generate DOCX"}), 500

# Entry point
if __name__ == "__main__":
    app.run(debug=True)